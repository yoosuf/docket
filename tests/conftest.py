from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.main import create_app
from app.modules.documents.rendering import DocxtplRenderer
from app.modules.documents.service import DocumentGenerationService, get_document_service
from app.modules.documents.storage import LocalFileStorage
from app.modules.documents.templates import FilesystemTemplateRepository


@pytest.fixture
def templates_dir(tmp_path: Path) -> Path:
    directory = tmp_path / "templates"
    directory.mkdir()

    doc = Document()
    doc.add_heading("SAMPLE", level=1)
    doc.add_paragraph("Hello {{ name }}")
    doc.save(directory / "sample_v1.docx")
    doc.save(directory / "sample_v2.docx")

    return directory


@pytest.fixture
def generated_dir(tmp_path: Path) -> Path:
    directory = tmp_path / "generated"
    directory.mkdir()
    return directory


@pytest.fixture
def storage(generated_dir: Path) -> LocalFileStorage:
    return LocalFileStorage(generated_dir)


class _UnavailableConverter:
    def convert_to_pdf(self, docx_bytes: bytes) -> bytes:
        raise AssertionError("PDF conversion should not be invoked in these tests")

    def convert_to_docx(self, doc_bytes: bytes) -> bytes:
        raise AssertionError("doc->docx conversion should not be invoked in these tests")


@pytest.fixture
def document_service(templates_dir: Path, storage: LocalFileStorage) -> DocumentGenerationService:
    return DocumentGenerationService(
        template_repository=FilesystemTemplateRepository(templates_dir),
        renderer=DocxtplRenderer(),
        converter=_UnavailableConverter(),
        storage=storage,
    )


@pytest.fixture
def client(document_service: DocumentGenerationService):
    app = create_app()
    app.dependency_overrides[get_document_service] = lambda: document_service
    with TestClient(app) as test_client:
        yield test_client
    app.dependency_overrides.clear()
