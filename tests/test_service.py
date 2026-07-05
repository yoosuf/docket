import io

import pytest
from docx import Document

from app.modules.documents.exceptions import TemplateNotFoundError
from app.modules.documents.rendering import DocxtplRenderer
from app.modules.documents.schemas import DocumentCreateRequest, OutputFormat
from app.modules.documents.service import DocumentGenerationService
from app.modules.documents.templates import FilesystemTemplateRepository


def _read_paragraphs(path):
    doc = Document(path)
    return [p.text for p in doc.paragraphs]


def _build_docx_bytes(placeholder_text: str) -> bytes:
    doc = Document()
    doc.add_paragraph(placeholder_text)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class _FakeDocConverter:
    """Fakes convert_to_docx by returning pre-built, valid docx bytes
    regardless of input. Real legacy .doc parsing can only be verified
    against genuine LibreOffice — this tests the orchestration (that a
    .doc template triggers doc->docx conversion before rendering, and
    docx->pdf conversion only when output_format is pdf), not LibreOffice
    itself.
    """

    def __init__(self, docx_bytes: bytes) -> None:
        self._docx_bytes = docx_bytes
        self.convert_to_docx_calls = 0
        self.convert_to_pdf_calls = 0

    def convert_to_docx(self, doc_bytes: bytes) -> bytes:
        self.convert_to_docx_calls += 1
        return self._docx_bytes

    def convert_to_pdf(self, docx_bytes: bytes) -> bytes:
        self.convert_to_pdf_calls += 1
        return b"%PDF-1.7 fake"


def test_generate_docx_renders_data_and_stores_file(document_service, generated_dir):
    request = DocumentCreateRequest(
        document_type="sample",
        data={"name": "Ada Lovelace"},
        output_format=OutputFormat.DOCX,
    )

    result = document_service.generate(request)

    assert result.output_format is OutputFormat.DOCX
    assert result.document_type == "sample"
    assert result.version == "v2"
    assert result.file_path.exists()
    assert result.file_path.parent == generated_dir
    assert result.size_bytes == result.file_path.stat().st_size
    assert result.file_path.name.endswith(f"_{result.document_id}.docx")
    assert "Hello Ada Lovelace" in _read_paragraphs(result.file_path)


def test_generate_with_explicit_version(document_service):
    request = DocumentCreateRequest(
        document_type="sample", data={"name": "Bob"}, version="v1"
    )

    result = document_service.generate(request)

    assert result.version == "v1"


def test_generate_unknown_document_type_raises(document_service):
    request = DocumentCreateRequest(document_type="unknown", data={})

    with pytest.raises(TemplateNotFoundError):
        document_service.generate(request)


def test_list_supported_document_types(document_service):
    descriptors = document_service.list_supported_document_types()

    doc_types = {d.document_type for d in descriptors}
    assert "sample" in doc_types


def test_get_document_returns_matching_metadata(document_service):
    created = document_service.generate(DocumentCreateRequest(document_type="sample", data={"name": "Ada"}))

    fetched = document_service.get_document(created.document_id)

    assert fetched is not None
    assert fetched.document_id == created.document_id
    assert fetched.document_type == "sample"
    assert fetched.size_bytes == created.size_bytes


def test_get_document_unknown_id_returns_none(document_service):
    assert document_service.get_document("0" * 32) is None


def test_list_documents_includes_generated_ones_newest_first(document_service):
    first = document_service.generate(DocumentCreateRequest(document_type="sample", data={"name": "A"}))
    second = document_service.generate(DocumentCreateRequest(document_type="sample", data={"name": "B"}))

    documents = document_service.list_documents()

    ids = [d.document_id for d in documents]
    assert first.document_id in ids
    assert second.document_id in ids
    assert ids.index(second.document_id) <= ids.index(first.document_id)


def test_delete_document_removes_file_and_reports_success(document_service):
    created = document_service.generate(DocumentCreateRequest(document_type="sample", data={"name": "Ada"}))
    assert created.file_path.exists()

    deleted = document_service.delete_document(created.document_id)

    assert deleted is True
    assert not created.file_path.exists()
    assert document_service.get_document(created.document_id) is None


def test_delete_document_unknown_id_returns_false(document_service):
    assert document_service.delete_document("0" * 32) is False


def test_generate_from_legacy_doc_template_converts_to_docx_before_rendering(templates_dir, storage):
    (templates_dir / "legacy_v1.doc").write_bytes(b"stand-in for a real binary .doc file")
    converter = _FakeDocConverter(_build_docx_bytes("Hello {{ name }}"))
    service = DocumentGenerationService(
        template_repository=FilesystemTemplateRepository(templates_dir),
        renderer=DocxtplRenderer(),
        converter=converter,
        storage=storage,
    )

    result = service.generate(DocumentCreateRequest(document_type="legacy", data={"name": "Ada"}))

    assert converter.convert_to_docx_calls == 1
    assert converter.convert_to_pdf_calls == 0
    assert result.file_path.suffix == ".docx"
    assert "Hello Ada" in _read_paragraphs(result.file_path)


def test_generate_pdf_from_legacy_doc_template_converts_both_directions(templates_dir, storage):
    (templates_dir / "legacy_v1.doc").write_bytes(b"stand-in for a real binary .doc file")
    converter = _FakeDocConverter(_build_docx_bytes("Hello {{ name }}"))
    service = DocumentGenerationService(
        template_repository=FilesystemTemplateRepository(templates_dir),
        renderer=DocxtplRenderer(),
        converter=converter,
        storage=storage,
    )

    result = service.generate(
        DocumentCreateRequest(document_type="legacy", data={"name": "Ada"}, output_format=OutputFormat.PDF)
    )

    assert converter.convert_to_docx_calls == 1
    assert converter.convert_to_pdf_calls == 1
    assert result.file_path.suffix == ".pdf"
